from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT = Path(__file__).with_name("Ripple XRPL Rights Clearing Agent Dossier.docx")


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def border(cell, color="D9D9D9"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:color"), color)


def cell_text(cell, text, bold=False, color=None, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Aptos"
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor(*color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    border(cell)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        c.width = Inches(widths[i])
        shade(c, "163A5F")
        cell_text(c, h, True, (255, 255, 255), 9)
    for row_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].width = Inches(widths[i])
            if row_idx % 2:
                shade(cells[i], "F3F7FA")
            cell_text(cells[i], value, False, None, 8.5)
    for row in table.rows:
        row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
    doc.add_paragraph().paragraph_format.space_after = Pt(3)
    return table


def add_para(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.12
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        r.bold = True
        r.font.size = Pt(10.5)
        rest = p.add_run(text[len(bold_lead):])
        rest.font.size = Pt(10.5)
    else:
        r = p.add_run(text)
        r.font.size = Pt(10.5)
    return p


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(12 if level == 1 else 7)
    p.paragraph_format.space_after = Pt(5)
    for r in p.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)
        r.font.name = "Aptos Display" if level == 1 else "Aptos"
    return p


def remove_paragraph_borders(paragraph):
    ppr = paragraph._p.get_or_add_pPr()
    borders = ppr.find(qn("w:pBdr"))
    if borders is not None:
        ppr.remove(borders)


def remove_style_borders(style):
    ppr = style._element.get_or_add_pPr()
    borders = ppr.find(qn("w:pBdr"))
    if borders is not None:
        ppr.remove(borders)


doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.7)
sec.bottom_margin = Inches(0.65)
sec.left_margin = Inches(0.72)
sec.right_margin = Inches(0.72)

styles = doc.styles
styles["Normal"].font.name = "Aptos"
styles["Normal"].font.size = Pt(10.5)
for style_name in ("Title", "Heading 1", "Heading 2"):
    styles[style_name].font.color.rgb = RGBColor(0, 0, 0)
remove_style_borders(styles["Title"])

title = doc.add_paragraph(style="Title")
remove_paragraph_borders(title)
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
title.paragraph_format.space_after = Pt(6)
title.add_run("Ripple XRPL Rights Clearing Agent Research Dossier")
sub = doc.add_paragraph()
sub.paragraph_format.space_after = Pt(16)
run = sub.add_run("Context before ideation  |  4 September 2026  |  Prepared for theFastandtheFungible")
run.italic = True
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(89, 89, 89)

heading(doc, "Executive answer")
add_para(doc, "The attractive problem is not simply making an AI buy an image. It is translating an agent’s natural-language commercial brief into a defensible decision that a particular asset, seller, licence and resulting artefact are usable for that brief. Payment is only one control point. The workflow must preserve the delegation, normalise rights terms, establish counterparty and asset integrity, bind the exact quote to payment, verify fulfilment and retain a reviewable evidence package.")
add_para(doc, "The direction fits the Ripple challenge because it creates a genuine commercial loop: a customer needs commercially usable content; an agent discovers offers and evaluates constraints; it purchases the selected licence; and it returns both a finished artefact and a machine-readable rights record. The non-negotiable design boundary is that a ledger transaction, a risk approval or a Content Credential does not itself prove that a licence is legally sufficient.")

heading(doc, "What the supplied material establishes")
add_table(doc, ["Question", "Evidence based answer"], [
    ("What is required?", "XRPL plus at least one successful XRPL transaction. The XRPL EVM Sidechain and other chains do not count."),
    ("What is recommended?", "The AI Starter Kit, x402 or MPP, and the repository resource skill. Use them when they make the commercial loop stronger; they are not hard requirements."),
    ("Is there a required dataset?", "No required business dataset was found in the supplied repository. It provides docs, examples and an optional resource skill."),
    ("What is the feedback hook?", "A builder-feedback telemetry mechanism connected to 10% of judging. It is not product data, an ideation input or an end-user feature."),
    ("What does a credible demo need?", "Need → discovery/comparison → bounded decision → XRPL settlement → useful delivery, plus clear safeguards and an observable receipt."),
], [1.55, 5.75])

heading(doc, "Interpreting the slides")
add_para(doc, "The supplied images are a t54 Trustline presentation, not a Ripple protocol specification. They frame Trustline as a transaction-level trust and risk layer for the agent economy: know the agent, confirm authorisation and intent, assess transaction risk, integrate across protocols and retain evidence for disputes. Treat the slides’ performance figures as vendor-presented claims, not independently verified benchmarks.")
add_table(doc, ["Failure pattern shown", "Rights-clearing interpretation"], [
    ("Prompt-injected path", "A retrieved page, tool response or embedded text changes the agent’s approved purchase route."),
    ("Counterfeit route", "The asset or seller looks genuine but the entity cannot grant the claimed licence."),
    ("Silent auto-renew", "A one-off licence quote hides a renewal, credit pack or later price commitment."),
    ("Redirected payment", "The merchant identity, invoice and XRPL destination do not resolve to the same approved counterparty."),
    ("Quiet downgrade", "The agent receives a lower or narrower licence than the one it selected."),
    ("Opinion laundering", "A selection is steered by poisoned reviews or a single low-integrity reputation source."),
], [1.8, 5.5])

heading(doc, "Ecosystem and technical roles")
heading(doc, "Ripple and XRPL", 2)
add_para(doc, "Ripple is the company; the XRP Ledger is the public ledger required by the challenge. The official materials position XRPL as the settlement and observability layer for agentic commerce. Relevant controls include testnet-first delivery, transaction simulation, explicit confirmation, transaction hashes, SourceTag, Memo, multisigning, DepositAuth and escrow. For this direction, the ledger should bind and settle an exact purchase—not hold private licence PDFs, customer briefs, personally identifiable information or full reasoning traces.")
heading(doc, "t54 not T45", 2)
add_para(doc, "The referenced company is t54.ai. It is not Ripple. Its core product, Trustline, is a pre-execution underwriting and evidence layer for agent-mediated financial actions. It takes transaction, agent, principal, policy, evidence and external-signal context and returns a risk posture. Its payment-facing products include x402-Secure and an XRPL x402 facilitator. This makes t54 relevant as a policy/evidence integration option, but it does not determine whether a third-party licence is legally adequate.")
heading(doc, "x402, rights metadata and provenance", 2)
add_para(doc, "x402 is an HTTP-native payment pattern: a seller returns a 402 response with payment requirements; the client supplies a payment payload; a facilitator verifies and settles; the seller returns the paid resource. It is a strong fit when a rights provider exposes a paid digital endpoint. ODRL and Creative Commons metadata can structure rights terms as machine-readable permissions, prohibitions, duties and constraints. C2PA Content Credentials can prove that provenance assertions are signed and tamper-evident. Neither x402 nor C2PA supplies legal interpretation or creates a licence entitlement on its own.")

heading(doc, "The rights-clearing decision flow")
flow = [
    ("1. Delegated brief", "Capture deliverable, commercial purpose, channel, territory, duration, budget, asset needs and exclusions in a versioned mandate."),
    ("2. Discovery", "Retrieve offers, terms, seller identity, asset provenance, price, expiry/renewal mechanics and delivery expectations."),
    ("3. Rights evaluation", "Use deterministic rules for permitted use, restrictions, attribution, derivative rights, exclusivity, AI restrictions and scope limits. Ambiguity goes to review."),
    ("4. Decision", "Rank only eligible offers. Record rejected alternatives. Create a payment intent bound to seller, asset/SKU, term version, amount, currency and expiry."),
    ("5. Pre-payment gate", "Check authority, policy, counterparty, quote integrity and total cost. Challenge missing or conflicting evidence before payment."),
    ("6. Settlement", "Execute the approved XRPL payment and reconcile the validated result with the expected amount, recipient and invoice."),
    ("7. Fulfilment", "Verify the returned asset/licence package, file hashes, version and seller receipt. Treat payment success and delivery success as separate states."),
    ("8. Evidence and reuse", "Return the artefact with a rights manifest; track expiry, renewal and later reuse as a new decision rather than an automatic right."),
]
add_table(doc, ["Stage", "Required decision or evidence"], flow, [1.45, 5.85])

doc.add_page_break()
heading(doc, "Pain point dossier")
add_table(doc, ["Where the flow fails", "Why it matters", "Minimum control"], [
    ("Brief to mandate", "“Commercially licensed” omits channels, geography, term, edits, attribution and exclusivity.", "Structured mandate; version and approval; unknown fields require review."),
    ("Terms semantics", "Labels such as “commercial” hide definitions, exceptions and combined-use restrictions.", "Machine-readable policy representation; legal escalation for ambiguity."),
    ("Seller and asset integrity", "A convincing page may lack authority to license a genuine-looking asset.", "Verified seller/payee relationship; origin capture; asset hash and provenance evidence."),
    ("Subscription and unit cost", "Cheap initial price can conceal renewal, credits or service limits.", "Term snapshot; total-cost cap; separate consent for recurring commitments."),
    ("Bounded authority", "A funded wallet is not a user mandate; agents can overbuy or select an ineligible bundle.", "Per-task payee, amount, currency, SKU and expiry limits enforced before signing."),
    ("Payment binding", "Payment details can change after selection or be replayed against another resource.", "Exact invoice/quote binding; nonce/expiry; idempotency; settlement receipt."),
    ("Settlement versus delivery", "A payment may validate while download, licence issuance or asset integrity fails.", "Two-stage state model; fulfilment receipt; reconciliation and escalation path."),
    ("Provenance versus legal proof", "A signed assertion proves who made it and that it was unmodified, not that it is legally complete.", "Show provenance, seller verification and policy eligibility as separate statuses."),
    ("Evidence and privacy", "Full traces may contain confidential briefs, personal data or proprietary strategy.", "Redacted off-chain evidence store; on-chain opaque correlation ID/hash only."),
], [1.55, 3.15, 2.6])

heading(doc, "XRPL implementation boundaries")
add_para(doc, "RLUSD is supported for direct XRPL payment, but receiver accounts need the issuer trust line; that onboarding and reserve requirement are real product friction. Do not promise RLUSD escrow. The current canonical RLUSD issuer configuration does not enable the trust-line locking prerequisite for native token escrow. XRPL escrow is useful for time/hash conditions, but it is not a content-delivery oracle or a dispute-resolution service.")
add_para(doc, "A validated payment is not enough by itself. Payment types that permit partial delivery must be reconciled using the validated delivered amount, not only a successful result code. Credentials, DIDs, NFTs, InvoiceID, SourceTag and Memos can help with attestation or correlation, but they do not natively express rights entitlement. Memos are public and constrained: use only an opaque correlation ID or hash.")
add_para(doc, "Multisigning and permission delegation can help constrain wallet operations, but they do not provide a general on-chain per-vendor/per-amount cap. A policy-aware signing or approval layer remains necessary to enforce the delegated mandate.")

heading(doc, "What an honest future prototype can claim")
add_para(doc, "The defensible claim is: the system records the terms and evidence it received, applies a disclosed policy, binds the approved payment to the selected offer, and explains why it paid. It should not claim to replace legal counsel, authenticate every creator, guarantee global rights clearance or make future downstream use lawful.")
add_para(doc, "The most meaningful test cases are rejection cases: editorial-only content proposed for an ad; a merchant whose payee changes after selection; a licence with a hidden renewal; a seller without adequate proof of authority; and a paid response whose delivered asset does not match the quoted package.")

doc.add_page_break()
heading(doc, "Sources and evidence notes")
sources = [
    "Official Ripple challenge README, Singhacks 2026, accessed 4 Sep 2026 — hard requirements, judging, resource status and feedback hook.",
    "XRPL Agentic Transactions and Agentic Payments with x402, XRPL Foundation, accessed 4 Sep 2026 — agent controls, x402 flow and T54 implementation context.",
    "Ripple XRPL AI Starter Kit announcement, 9 Jun 2026 — Ripple positioning and starter-kit context.",
    "Trustline Overview and Risk Engine, t54, accessed 4 Sep 2026 — Trustline product scope, inputs, decisions and evidence model.",
    "XRPL x402 Facilitator Overview, t54, accessed 4 Sep 2026 — HTTP 402 request-payment-delivery flow using XRP/RLUSD.",
    "x402 Specification v1, x402 Foundation, accessed 4 Sep 2026 — protocol roles and the explicit exclusion of client budget management.",
    "ODRL Information Model 2.2, W3C; Creative Commons licensing metadata FAQ — machine-readable rights representation and discoverability.",
    "C2PA Content Credentials Specification and Guiding Principles — signed, tamper-evident provenance and its non-judgment limitation.",
    "RLUSD on the XRP Ledger, Reserves, Escrow, Partial Payments, Credentials and Permission Delegation, Ripple/XRPL Foundation, accessed 4 Sep 2026 — trust lines, reserve, escrow, reconciliation, credential and signing limits.",
    "US Copyright Office Copyright and Artificial Intelligence materials — background on continuing AI/copyright uncertainty; not a jurisdictional legal conclusion.",
]
for idx, source in enumerate(sources, 1):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"{idx}. {source}")
    r.font.size = Pt(8.8)

heading(doc, "Research limitations")
add_para(doc, "The slides’ performance and partnership figures are vendor claims, not independently benchmarked results. No actual supplier contract, licensing API, term corpus, campaign jurisdiction or legal review was provided. This dossier identifies the evidence and decision boundaries; it does not decide the legal sufficiency of any actual asset. The evidence is sufficient to move into solution framing once the team selects a first asset category and jurisdiction.")

for section in doc.sections:
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("Ripple XRPL Rights Clearing Agent Research Dossier  |  4 September 2026")
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(100, 100, 100)

doc.save(OUT)
print(OUT)
